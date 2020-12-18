#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import os,sqlite3,time,shutil
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.shared import RGBColor
from urllib.parse import urlparse
from lib.common.utils import Utils
from lib.TestProxy import testProxy
from lib.Database import DatabaseType
from lib.common.CreatLog import creatLog
from lib.reports.creat_api import Creat_api
from lib.common.cmdline import CommandLines
from lib.reports.creat_tree import Creat_tree
from lib.reports.creat_suggest import Creat_suggest
from lib.reports.creat_vuln_detail import Creat_vuln_detail


class Docx_replace():

    def __init__(self,projectTag):
        docLang = Utils().getMyWord("{lang}")
        self.projectTag = projectTag
        self.tmp_filepath = "doc" + os.sep + "template" + os.sep + docLang + ".docx"
        self.new_filepath = "reports" + os.sep + "tmp_" + self.projectTag + ".docx"
        self.log = creatLog().get_logger()

    def vuln_judge(self):
        vuln_judges = []
        vuln_h_num = [0,6]
        vuln_m_num = [0,2]
        vuln_l_num = [0,1]
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select * from vuln"
        cursor.execute(sql)
        vuln_infos = cursor.fetchall()
        for vuln_info in vuln_infos:
            if vuln_info[3] == "INFO":
                vuln_m_num[0] = vuln_m_num[0] + 1
            if vuln_info[3] == "unAuth" and vuln_info[4] == 1:
                vuln_m_num[0] = vuln_m_num[0] + 1
            if vuln_info[3] == "CORS":
                vuln_l_num[0] = vuln_l_num[0] + 1
            if vuln_info[3] == "unAuth" and vuln_info[4] == 2:
                vuln_l_num[0] = vuln_l_num[0] + 1
            if vuln_info[3] == "passWord":
                vuln_h_num[0] = vuln_h_num[0] + 1
            if vuln_info[3] == "BAC":
                vuln_m_num[0] = vuln_m_num[0] + 1
            if vuln_info[3] == "upLoad":
                vuln_h_num[0] = vuln_h_num[0] + 1

        vuln_score =  vuln_h_num[0]*vuln_h_num[1] + vuln_m_num[0]*vuln_m_num[1] + vuln_l_num[0]*vuln_l_num[1]
        vuln_judges.append(vuln_score)
        vuln_judges.append(vuln_h_num[0])
        vuln_judges.append(vuln_m_num[0])
        vuln_judges.append(vuln_l_num[0])
        return vuln_judges


    def docxReplace(self, document):
        cmd = CommandLines().cmd()
        ipAddr = testProxy(cmd,0)
        end_time = time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())
        report_time = time.strftime("%Y-%m-%d %H:%M", time.localtime())
        main_url = DatabaseType(self.projectTag).getURLfromDB()
        parse_url = urlparse(main_url)
        host = parse_url.netloc
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select id from js_file"
        cursor.execute(sql)
        js_all_files = cursor.fetchall()
        js_num = len(js_all_files)
        sql = "select path from js_file"
        cursor.execute(sql)
        jsfilelist = cursor.fetchall()
        js_paths=''
        for js in jsfilelist:
            jspath = "◆ " + js[0] + "\n"
            js_paths = js_paths + jspath
        sql ="select id from api_tree where success = 1 or success = 2"
        cursor.execute(sql)
        api_list = cursor.fetchall()
        api_num = len(api_list)
        vuln_infos = Docx_replace(self.projectTag).vuln_judge()
        vuln_h_num = vuln_infos[1]
        vuln_m_num = vuln_infos[2]
        vuln_l_num = vuln_infos[3]
        vuln_num = vuln_h_num + vuln_m_num + vuln_l_num
        vuln_score = vuln_infos[0]
        if vuln_score >= 18:
            sec_lv = Utils().getMyWord("{risk_h}")
        elif vuln_score < 18 and vuln_score >= 10:
            sec_lv = Utils().getMyWord("{risk_m}")
        elif vuln_score < 10 and vuln_score >= 5:
            sec_lv = Utils().getMyWord("{risk_l}")
        else:
            sec_lv = Utils().getMyWord("{risk_n}")
        sql = "select vaule from info where name='time'"
        cursor.execute(sql)
        time_in_info = cursor.fetchone()
        timeArray = time.localtime(int(time_in_info[0]))
        start_time = time.strftime("%Y-%m-%d %H:%M:%S", timeArray)
        type = CommandLines().cmd().type
        if type == "simple":
            scan_type = Utils().getMyWord("{mode_simple}")
        else:
            scan_type = Utils().getMyWord("{mode_adv}")
        scan_min = int(end_time.split(":")[-2]) - int(start_time.split(":")[-2])
        if int(scan_min) >= 1:
            end_time_one = int(end_time.split(":")[-1]) + int(scan_min) * 60
            scan_time = int(end_time_one) - int(start_time.split(":")[-1])
        else:
            scan_time = int(end_time.split(":")[-1]) - int(start_time.split(":")[-1])
        vuln_list = ''
        sql = "select id from vuln where type='unAuth'"
        cursor.execute(sql)
        num_auth = cursor.fetchall()
        if len(num_auth) != 0:
            vuln_list =  vuln_list +  "◆ " + Utils().getMyWord("{vuln_unauth_num}") + str(len(num_auth)) + Utils().getMyWord("{ge}") + "\n"
        sql = "select id from vuln where type='CORS'"
        cursor.execute(sql)
        num_cors = cursor.fetchall()
        if len(num_cors) != 0:
            vuln_list = vuln_list + "◆ " +  Utils().getMyWord("{vuln_cors_num}") + str(len(num_cors)) + Utils().getMyWord("{ge}") + "\n"
        sql = "select id from vuln where type='INFO'"
        cursor.execute(sql)
        num_info = cursor.fetchall()
        if len(num_info) != 0:
            vuln_list = vuln_list +  "◆ " + Utils().getMyWord("{vuln_info_num}") + str(len(num_info)) + Utils().getMyWord("{ge}") + "\n"
        sql = "select id from vuln where type='passWord'"
        cursor.execute(sql)
        num_passWord = cursor.fetchall()
        if len(num_passWord) != 0:
            vuln_list = vuln_list + "◆ " + Utils().getMyWord("{vuln_passWord_num}") + str(len(num_passWord)) + Utils().getMyWord("{ge}") + "\n"
        sql = "select id from vuln where type='BAC'"
        cursor.execute(sql)
        num_BAC = cursor.fetchall()
        if len(num_BAC) != 0:
            vuln_list = vuln_list + "◆ " + Utils().getMyWord("{vuln_BAC_num}") + str(len(num_BAC)) + Utils().getMyWord("{ge}") + "\n"
        sql = "select id from vuln where type='upLoad'"
        cursor.execute(sql)
        num_upload = cursor.fetchall()
        if len(num_upload) != 0:
            vuln_list = vuln_list + "◆ " + Utils().getMyWord("{vuln_upload_num}") + str(
                len(num_upload)) + Utils().getMyWord(
                "{ge}") + "\n"
        sql = "select id from vuln where type='SQL'"
        cursor.execute(sql)
        num_sql = cursor.fetchall()
        if len(num_sql) != 0:
            vuln_list = vuln_list + "◆ " + Utils().getMyWord("{vuln_sql_num}") + str(
                len(num_sql)) + Utils().getMyWord(
                "{ge}") + "\n"
        cookies = CommandLines().cmd().cookie
        if cookies:
            extra_cookies = cookies
        else:
            extra_cookies = Utils().getMyWord("{no_extra_cookies}")
        head = CommandLines().cmd().head
        if head != "Cache-Control:no-cache":
             extra_head = head
        else:
             extra_head = Utils().getMyWord("{no_extra_head}")
        try:
            DICT = {
                "{report_number}": "PF-API-" + self.projectTag,
                "{report_date}": "%s" % (report_time),
                "{target_host}": "%s" % (host),
                "{target_url}": "%s" % (main_url),
                "{js_num}": "%s" % (js_num),
                "{start_time}": "%s" % (start_time),
                "{scan_time}" :"%s" % (scan_time),
                "{scan_type}": "%s" % (scan_type),
                "{api_num}": "%s" % (api_num),
                "{vuln_num}": "%s" % (vuln_num),
                "{vuln_h_num}": "%s" % (vuln_h_num),
                "{vuln_m_num}": "%s" % (vuln_m_num),
                "{vuln_l_num}": "%s" % (vuln_l_num),
                "{unauth_vuln}": "%s" % ("unauth_vuln"),
                "{vuln_list}": "%s" % (vuln_list),
                "{scan_ip}": "%s" % (ipAddr),
                "{extra_cookies}": "%s" % (extra_cookies),
                "{extra_head}": "%s" % (extra_head)
            }
            self.log.debug("word—report正常替换")
        except Exception as e:
            self.log.error("[Err] %s" % e)

        for table in document.tables:
            for row in range(len(table.rows)):
                for col in range(len(table.columns)):
                    for key, value in DICT.items():
                        if key in table.cell(row, col).text:
                            table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

        for para in document.paragraphs:
            for i in range(len(para.runs)):
                for key, value in DICT.items():
                    if key in para.runs[i].text:
                        para.runs[i].text = para.runs[i].text.replace(key, value)

        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{js_list}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace("{js_list}", "%s" % (js_paths))
                    para.runs[i].font.size = Pt(10)
                    para.runs[i].font.name = "Arial"

        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{end_time}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace("{end_time}", "%s" % (end_time))

        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{sec_lv}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace("{sec_lv}", "%s" % (sec_lv))
                    para.runs[i].font.size = Pt(14)
                    if sec_lv == Utils().getMyWord("{risk_n}"):
                        para.runs[i].font.color.rgb = RGBColor(139,137,137)
                    elif sec_lv == Utils().getMyWord("{risk_l}"):
                        para.runs[i].font.color.rgb = RGBColor(46, 139,87)
                    elif sec_lv == Utils().getMyWord("{risk_m}"):
                        para.runs[i].font.color.rgb = RGBColor(205, 55, 0)
                    elif sec_lv == Utils().getMyWord("{risk_h}"):
                        para.runs[i].font.color.rgb = RGBColor(238, 0, 0)

        try:
            Creat_vuln_detail(self.projectTag).creat_detail(document)
            self.log.debug("正确获取vuln_detail替换内容")
        except Exception as e:
            self.log.error("[Err] %s" % e)
        try:
            Creat_api(self.projectTag).creat_api(document)
            self.log.debug("正确获取api替换内容")
        except Exception as e:
            self.log.error("[Err] %s" % e)
        try:
            Creat_suggest(self.projectTag).creat_suggest(document)
            self.log.debug("正确获取suggest替换内容")
        except Exception as e:
            self.log.error("[Err] %s" % e)
        try:
            Creat_tree(self.projectTag).tree_list(document)
            self.log.debug("正确获取extra替换内容")
        except Exception as e:
            self.log.error("[Err] %s" % e)

        return document


    def mainReplace(self):
        document = Document(self.tmp_filepath)
        document = Docx_replace(self.projectTag).docxReplace(document)
        document.save(self.new_filepath)

    def docMove(self,nameDoc):
        shutil.copyfile(self.new_filepath, nameDoc)

    def docDel(self):
        os.remove(self.new_filepath)
