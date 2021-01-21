#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import os,sqlite3,time,json
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from urllib.parse import urlparse
from lib.common.utils import Utils
from lib.Database import DatabaseType
from lib.common.CreatLog import creatLog
from lib.common.cmdline import CommandLines

class Creat_tree():
    def __init__(self,projectTag):
        self.projectTag = projectTag
        self.treeStr = ""

    def locat_suggest(self, document):
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "extra" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace('extra', '')
                    para1 = para.insert_paragraph_before("")
        return para1

    def move_table_after(self, table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    def creat_table(self, document, vuln_info, para2):
        colorStr = 'F5F5F5'
        row = col = 1
        table = document.add_table(row, col)
        table.rows[0].cells[0].text = "%s" % (vuln_info)
        para3 = para2.insert_paragraph_before("")
        Creat_tree(self.projectTag).tabBgColor(table, col, colorStr)
        Creat_tree(self.projectTag).move_table_after(table, para2)

    def tabBgColor(self,table, cols, colorStr):
        shading_list = locals()
        for i in range(cols):
            shading_list['shading_elm_' + str(i)] = parse_xml(
                r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor=colorStr))
            table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_' + str(i)])

    def tree_list(self, document):
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select path from api_tree where success = 1 or success = 2"
        cursor.execute(sql)
        apinfo = cursor.fetchall()
        sql = "select path from js_file"
        cursor.execute(sql)
        jsInfo = cursor.fetchall()
        jsInfo = jsInfo + apinfo
        list1 = [[] for i in range(len(jsInfo))]
        k = 0
        for js in jsInfo:
            js = js[0]
            path = js.split("/")[3:]
            list1[k].append(path)
            k = k + 1
        for i in range(len(list1)):
            for j in range(len(list1[i][0])):
                if not j == 0:
                    list1[i][0][j] = list1[i][0][j - 1] + "-->" * j + list1[i][0][j]
        list_max = []
        for i in range(len(list1)):
            for j in range(len(list1[i])):
                maxlen = len(list1[i][j])
                list_max.append(maxlen)
        maxlen = max(list_max)
        level_list = [[] for i in range(maxlen)]
        for i in range(len(list1)):
            for j in range(maxlen):
                if list1[i][0][j:j + 1] != []:
                    level_list[j].append(list1[i][0][j:j + 1][0])
        new_level_list = []
        for i in level_list:
            param = list(set(i))
            new_level_list.append(param)
        level_add = 1
        main_url = DatabaseType(self.projectTag).getURLfromDB()
        parse_url = urlparse(main_url)
        host = parse_url.netloc
        self.treeStr = self.treeStr + host + "\n"
        for trunk in new_level_list[0]:
            self.treeStr = self.treeStr + ("|----" + trunk) + "\n"
            self.loop_branch(new_level_list, level_add, trunk)
        para1 = Creat_tree(self.projectTag).locat_suggest(document)
        para2 = para1.insert_paragraph_before("")
        run2 = para2.add_run("当前网站资源树如下：")
        run2.font.name = "Arial"
        run2.font.size = Pt(11)
        run2.font.bold = True
        Creat_tree(self.projectTag).creat_table(document, self.treeStr, para2)
        #
        #
        # self.treeStr = ''
        # sql = "select path from api_tree"
        # cursor.execute(sql)
        # jsInfo = cursor.fetchall()
        # list1 = [[] for i in range(len(jsInfo))]
        # k = 0
        # for js in jsInfo:
        #     js = js[0]
        #     path = js.split("/")[3:]
        #     list1[k].append(path)
        #     k = k + 1
        # for i in range(len(list1)):
        #     for j in range(len(list1[i][0])):
        #         if not j == 0:
        #             list1[i][0][j] = list1[i][0][j - 1] + "-->" * j + list1[i][0][j]
        # list_max = []
        # for i in range(len(list1)):
        #     for j in range(len(list1[i])):
        #         maxlen = len(list1[i][j])
        #         list_max.append(maxlen)
        # maxlen = max(list_max)
        # level_list = [[] for i in range(maxlen)]
        # for i in range(len(list1)):
        #     for j in range(maxlen):
        #         if list1[i][0][j:j + 1] != []:
        #             level_list[j].append(list1[i][0][j:j + 1][0])
        # new_level_list = []
        # for i in level_list:
        #     param = list(set(i))
        #     new_level_list.append(param)
        # level_add = 1
        # for trunk in new_level_list[0]:
        #     self.treeStr = self.treeStr + ("|----" + trunk) + "\n"
        #     self.loop_branch(new_level_list, level_add, trunk)
        # para2.add_run("\n\n" + "api" + "\n" + self.treeStr)


    def loop_branch(self,new_level_list, level_add, trunk):
        if level_add == len(new_level_list):
            level_add = 1
        for branch in new_level_list[level_add]:
            if branch.split("-->" * (level_add))[-2] == trunk:
                self.treeStr = self.treeStr + ("     " * (level_add) + "|----" + branch.split("-->")[-1]) + "\n"
                self.loop_branch(new_level_list, level_add + 1, branch)
            else:
                continue

