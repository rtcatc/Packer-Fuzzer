#!/usr/bin/env python3
# -*- encoding: utf-8 -*-

import os,sqlite3,time,json
from docx import Document   #用来建立一个word对象
from docx.shared import Pt  #用来设置字体的大小
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from urllib.parse import urlparse
from lib.common.utils import Utils
from lib.Database import DatabaseType
from lib.common.cmdline import CommandLines


class Creat_suggest():

    def __init__(self,projectTag):
        self.projectTag = projectTag
        self.creat_num = 1

    def locat_suggest(self, document):
        for para in document.paragraphs:
            for i in range(len(para.runs)):
                if "{suggest_foryou}" in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace('{suggest_foryou}', '')
                    para1 = para.insert_paragraph_before("")
        return para1

    def creat_suggest(self,document):
        para1 = Creat_suggest(self.projectTag).locat_suggest(document)
        projectDBPath = DatabaseType(self.projectTag).getPathfromDB() + self.projectTag + ".db"
        connect = sqlite3.connect(os.sep.join(projectDBPath.split('/')))
        cursor = connect.cursor()
        connect.isolation_level = None
        sql = "select * from vuln"
        cursor.execute(sql)
        vuln_infos = cursor.fetchall()
        flag1 = flag2 = flag3 = flag4 = flag5 = flag6= flag7 = 0
        for vuln_info in vuln_infos:
            if vuln_info[3] == "unAuth":
                flag1 = 1
            elif vuln_info[3] == "INFO":
                flag2 = 1
            elif vuln_info[3] == "CORS":
                flag3 = 1
            elif vuln_info[3] == "SQL":
                flag4 = 1
            elif vuln_info[3] == "upLoad":
                flag5 = 1
            elif vuln_info[3] == "passWord":
                flag6 = 1
            elif vuln_info[3] == "BAC":
                flag7 = 1
        para2 = para1.insert_paragraph_before("")
        if flag1 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_unauth_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_unauth_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_unauth_3}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag2 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_info_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_info_2}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag3 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_cors_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_cors_2}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag4 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_sqli_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_sqli_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_sqli_3}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag5 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_upload_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_upload_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_upload_2}") + "\n" + Utils().getMyWord("{r_sug_upload_3}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag6 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_password_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_password_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_password_3}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        if flag7 == 1:
            run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_bac_1}") + "\n")
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_bac_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_bac_3}") + "\n")
            run2.font.name = "Arial"
            run2.font.size = Pt(10)
            self.creat_num = self.creat_num + 1
        run = para2.add_run("4." + str(self.creat_num) + Utils().getMyWord("{r_sug_g_1}") + "\n")
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.bold = True
        run2 = para2.add_run("◆ " + Utils().getMyWord("{r_sug_g_2}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_g_3}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_g_4}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_g_5}") + "\n" + "◆ " + Utils().getMyWord("{r_sug_g_6}") + "\n")
        run2.font.name = "Arial"
        run2.font.size = Pt(10)
